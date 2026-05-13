"""在毕业论文 Word 文档中插入 4.2.1、4.3.1、4.3.2 小节内容。"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/mnt/d/Download/毕业论文_扩展版_20260512.docx"
DST = "/mnt/d/Download/毕业论文_扩展版_20260512_revised.docx"


def get_pPr(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._p.insert(0, pPr)
    return pPr


def _set_fixed_line(para, twentieths=440):
    pPr = get_pPr(para)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(twentieths))
    spacing.set(qn("w:lineRule"), "exact")


def insert_para_after(ref_para, doc):
    new_p = OxmlElement("w:p")
    ref_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, doc._body)


def insert_para_before(ref_para, doc):
    new_p = OxmlElement("w:p")
    ref_para._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, doc._body)


def set_heading3(para, text):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_fixed_line(para, 440)


def set_body(para, text, bold_prefix=None):
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(12)
    r2 = para.add_run(text)
    r2.bold = False
    r2.font.size = Pt(12)
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(24)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_fixed_line(para, 440)


def find_para_by_prefix(doc, prefix):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i, p
    return None, None


# ----------------------------------------------------------------
# 段落内容（不含弯引号，避免 Python 词法歧义）
# ----------------------------------------------------------------

CONTENT_421_HEADING = "4.2.1　MinerU解析质量增强策略"

CONTENT_421_INTRO = (
    "MinerU在工艺规程PDF解析中存在表格截断和静默失败两类问题，"
    "本文针对这两类问题分别设计了修复策略。"
)

CONTENT_421_P1_BOLD = "表格内容截断与标题丢失。"
CONTENT_421_P1 = (
    "MinerU在处理跨页表格时存在两类典型问题。"
    "其一，HTML table_body字段完整但text字段被截断：通过比对<tr>标签数与Markdown数据行数，"
    "当HTML行数多于Markdown行数时回退使用HTML并转换为Markdown格式，确保表格内容完整。"
    "其二，跨页表格的标题块在合并时可能被消耗：引入pending_caption机制，"
    "在四类合并场景（C1-C4）下统一保存标题并向前传递，修复了表9-2、表12-10等标题丢失问题。"
)

CONTENT_421_P2_BOLD = "解析静默失败。"
CONTENT_421_P2 = (
    "针对上述问题，本文通过扩展bbox渲染区域（向上扩展60pt）并调用Qwen-VL-max视觉语言模型"
    "对裁剪图像进行表格内容识别，将提取结果回填至对应JSON块的text字段，"
    "共修复3处静默失败的空表格块。提取结果经三级后处理过滤："
    "首先按空行分段取首段，再检测列数突变（>=3列），"
    "最后通过首列数字/非数字切换识别下一张表的表头，防止VLM输出越界捕获相邻表格内容。"
)

CONTENT_431_HEADING = "4.3.1　语义感知切分策略设计"

CONTENT_431_P1 = (
    "针对工艺规程的文档特点，本文设计了两阶段语义感知切分策略。"
    "第一阶段对拼接文本进行标题注入预处理：通过正则表达式识别形如"
    "\"7.1.1 安全基本概念\"的编号标题（单行、长度<=60字符），"
    "分别注入##/###/####前缀，使MarkdownHeaderTextSplitter"
    "能以章节为强边界、以节为弱边界进行切分，彻底消除跨节内容混入问题。"
    "第二阶段在各节内部采用贪心积累策略（_greedy_accumulate）：H1/H2级标题触发立即刷新，"
    "H3/H4级标题触发贪心积累——仅当积累长度超过chunk_size时才切分，且切分点严格在节内，"
    "不产生跨节重叠。"
)

CONTENT_431_P2_BOLD = "表格语义完整切分。"
CONTENT_431_P2 = (
    "在表格切分上，采用四种场景的跨页合并规则（C1-C4），并引入"
    "\"挂起标题（pending_caption）\"机制：当合并操作导致表格标题块被消耗时，"
    "该标题暂存并向前传递给下一个独立表格块，确保每张表格的语义描述不丢失。"
    "最终将文本块与表格块交由同一切分流水线统一处理，输出结构一致的Document列表。"
)

CONTENT_431_P3 = (
    "经上述策略处理，实验语料由原始92个切分单元增至106个，chunk中位长度833字符，"
    "各切分单元均在章节边界内完整，未出现跨节内容混合。"
)

CONTENT_432_HEADING = "4.3.2　分割参数消融实验"


def main():
    doc = Document(SRC)

    i227, p227 = find_para_by_prefix(doc, "图4-1以工艺规程第9章")
    i228, p228 = find_para_by_prefix(doc, "4.3  文本分割参数实验")
    i229, p229 = find_para_by_prefix(doc, "文本分割参数（chunk_size与chunk_overlap）")

    print(f"p227({i227}): {p227.text[:50]}")
    print(f"p228({i228}): {p228.text[:50]}")
    print(f"p229({i229}): {p229.text[:50]}")

    # 从后往前插入，避免索引漂移
    # --- 4.3.2 标题（在 p229 前）---
    h_432 = insert_para_before(p229, doc)
    set_heading3(h_432, CONTENT_432_HEADING)

    # --- 4.3.1 小节（在 p228 后，倒序插入保证顺序正确）---
    b431_3 = insert_para_after(p228, doc)
    set_body(b431_3, CONTENT_431_P3)

    b431_2 = insert_para_after(p228, doc)
    set_body(b431_2, CONTENT_431_P2, bold_prefix=CONTENT_431_P2_BOLD)

    b431_1 = insert_para_after(p228, doc)
    set_body(b431_1, CONTENT_431_P1)

    h_431 = insert_para_after(p228, doc)
    set_heading3(h_431, CONTENT_431_HEADING)

    # --- 4.2.1 小节（在 p227 后，倒序插入）---
    b421_2 = insert_para_after(p227, doc)
    set_body(b421_2, CONTENT_421_P2, bold_prefix=CONTENT_421_P2_BOLD)

    b421_1 = insert_para_after(p227, doc)
    set_body(b421_1, CONTENT_421_P1, bold_prefix=CONTENT_421_P1_BOLD)

    b421_intro = insert_para_after(p227, doc)
    set_body(b421_intro, CONTENT_421_INTRO)

    h_421 = insert_para_after(p227, doc)
    set_heading3(h_421, CONTENT_421_HEADING)

    doc.save(DST)
    print(f"\n保存完成: {DST}")

    # 验证
    doc2 = Document(DST)
    count = sum(1 for p in doc2.paragraphs if "4.2.1" in p.text or "4.3.1" in p.text or "4.3.2" in p.text)
    print(f"验证：包含新标题的段落数 = {count}（应为3）")


if __name__ == "__main__":
    main()
