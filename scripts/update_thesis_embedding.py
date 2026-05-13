"""更新论文中 embedding 模型相关描述：bge-large-zh-v1.5 → bge-m3 为主力模型。"""

from docx import Document
from docx.oxml.ns import qn
import copy

SRC = "/mnt/d/Download/毕业论文_扩展版_20260512_revised.docx"
DST = "/mnt/d/Download/毕业论文_扩展版_20260512_revised.docx"


def replace_runs_text(para, old: str, new: str):
    """在段落的 runs 里做文本替换（保留格式）。"""
    full = para.text
    if old not in full:
        return False
    # 简单替换：把所有 run 文本拼起来替换，再写回第一个 run
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = full.replace(old, new)
    else:
        para.add_run(full.replace(old, new))
    return True


def find_para(doc, prefix: str):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i, p
    return None, None


def set_para_text(para, text: str):
    """清空段落内容，重新写入文本（保留第一个 run 的格式）。"""
    fmt_run = para.runs[0] if para.runs else None
    for run in para.runs:
        run.text = ""
    if fmt_run:
        fmt_run.text = text
    else:
        para.add_run(text)


def main():
    doc = Document(SRC)

    # ── [84] 粗召回阶段 ──────────────────────────────────────────────────────
    _, p84 = find_para(doc, "基于上述特性，本文采用两阶段检索架构")
    if p84:
        replace_runs_text(
            p84,
            "BGE-Large-ZH-v1.5（双编码器）",
            "BAAI/bge-m3（双编码器）",
        )
        print("[84] 更新粗召回阶段 embedding 描述")

    # ── [108] 离线索引阶段描述 ───────────────────────────────────────────────
    _, p108 = find_para(doc, "离线索引阶段：原始PDF文档经MinerU框架")
    if p108:
        replace_runs_text(p108, "由BGE嵌入模型转化为1024维向量", "由bge-m3嵌入模型转化为1024维向量")
        print("[108] 更新离线索引阶段描述")

    # ── [152-153] 3.4.1 BGE嵌入模型 ─────────────────────────────────────────
    i152, p152 = find_para(doc, "3.4.1")
    if p152:
        p153 = doc.paragraphs[i152 + 1]
        new_153 = (
            "本文主要实验选用BAAI/bge-m3作为嵌入模型，选择理由有二："
            "其一，bge-m3支持超过100种语言，具备跨语言对齐能力，可直接处理中英混合工艺文档及跨语言检索场景；"
            "其二，bge-m3的最大输入长度达8192 token，远超传统BERT系模型的512 token上限，"
            "使chunk_size在256至1024 token范围内均可完整嵌入，避免了因截断导致的向量语义缺失问题。"
            "该模型基于XLM-RoBERTa架构，以大规模多语言语料进行对比学习预训练，"
            "输出向量维度1024，在中文语义检索C-MTEB基准测试中性能优于同系列的bge-large-zh-v1.5。"
            "模型在本地RTX 3070 Ti Laptop GPU上运行，显存占用约2.3GB，推理延迟约0.6秒/批（32条）。"
        )
        set_para_text(p153, new_153)
        print("[153] 更新 3.4.1 主体描述")

    # ── [155] Chroma 向量库描述 ──────────────────────────────────────────────
    _, p155 = find_para(doc, "向量库以collection为单位进行组织")
    if p155:
        replace_runs_text(p155, "经BGE嵌入后存入Chroma", "经bge-m3嵌入后存入Chroma")
        print("[155] 更新 Chroma 描述")

    # ── [217] 实验环境知识库描述 ─────────────────────────────────────────────
    _, p217 = find_para(doc, "向量知识库：hydro_manual")
    if p217:
        replace_runs_text(p217, "使用BAAI/bge-large-zh-v1.5嵌入", "使用BAAI/bge-m3嵌入")
        print("[217] 更新实验环境知识库描述")

    # ── [244] 4.4 嵌入模型对比实验引言 ──────────────────────────────────────
    _, p244 = find_para(doc, "为评估不同嵌入模型在工业规程专业文档上的检索性能")
    if p244:
        new_244 = (
            "为评估不同嵌入模型在工业规程专业文档上的检索性能，本文以BAAI/bge-m3为主力模型，"
            "并与bge-large-zh-v1.5、bge-small-zh-v1.5、moka-ai/m3e-base、"
            "text2vec-base-chinese四种模型进行对比，以Hit@3、MRR及推理延迟为评估指标。"
            "bge-m3的选用依据其两项关键优势：支持超过100种语言的跨语言嵌入能力，"
            "以及8192 token的长上下文支持，可在不截断的前提下嵌入更大粒度的文本块。"
        )
        set_para_text(p244, new_244)
        print("[244] 更新 4.4 嵌入模型对比实验描述")

    doc.save(DST)
    print(f"\n保存完成: {DST}")

    # 验证
    doc2 = Document(DST)
    _, v = find_para(doc2, "本文主要实验选用BAAI/bge-m3")
    print(f"验证: {'✓ bge-m3 写入成功' if v else '✗ 未找到更新内容'}")


if __name__ == "__main__":
    main()
